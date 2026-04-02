#!/usr/bin/env python3
"""Generate Ohanafy-branded reference.docx for pandoc --reference-doc usage.

This creates a DOCX with custom styles that pandoc will pick up when converting
Markdown to DOCX. Run once to regenerate the template after brand changes.

Usage: python3 scripts/generate-reference-docx.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT = os.path.join(PROJECT_ROOT, "docs", "templates", "reference.docx")

# Brand colors
TRUE_BLACK = RGBColor(0x00, 0x00, 0x00)
TRUE_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MELLOW = RGBColor(0xE4, 0xF2, 0x23)
CORK = RGBColor(0xF4, 0xF2, 0xF0)
DARK_DENIM = RGBColor(0x4A, 0x5F, 0x80)
DARK_GREY = RGBColor(0x54, 0x54, 0x54)

# Font (Geist with fallbacks pandoc/Word can use)
FONT_NAME = "Geist"
FONT_FALLBACK = "Calibri"


def set_font(run, name=FONT_NAME, size=None, bold=None, color=None):
    """Apply font settings to a run."""
    run.font.name = name
    # Set fallback for systems without Geist
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), FONT_FALLBACK)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color


def style_paragraph(style, font_size, bold=False, color=TRUE_BLACK, space_before=0, space_after=6, alignment=None):
    """Configure a paragraph style."""
    fmt = style.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if alignment:
        fmt.alignment = alignment
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(font_size)
    font.bold = bold
    font.color.rgb = color
    # Set east-asian / complex script fallback
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_FALLBACK)


def main():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Styles ---
    # Title (54pt Bold)
    style_paragraph(doc.styles["Title"], font_size=54, bold=True, color=TRUE_BLACK, space_after=12)

    # Subtitle (36pt Thin)
    style_paragraph(doc.styles["Subtitle"], font_size=36, bold=False, color=DARK_DENIM, space_after=12)

    # Heading 1 (24pt Regular)
    style_paragraph(doc.styles["Heading 1"], font_size=24, bold=False, color=TRUE_BLACK, space_before=18, space_after=8)

    # Heading 2 (18pt Bold)
    style_paragraph(doc.styles["Heading 2"], font_size=18, bold=True, color=DARK_DENIM, space_before=14, space_after=6)

    # Heading 3 (16pt Bold)
    style_paragraph(doc.styles["Heading 3"], font_size=16, bold=True, color=DARK_DENIM, space_before=12, space_after=4)

    # Normal / Body (16pt Regular)
    style_paragraph(doc.styles["Normal"], font_size=16, bold=False, color=TRUE_BLACK, space_after=6)

    # --- Header with branded tagline ---
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = hp.add_run("Modern beverage runs on ")
    set_font(run1, size=Pt(14), bold=True, color=DARK_DENIM)
    run2 = hp.add_run("OHANAFY.")
    set_font(run2, size=Pt(14), bold=True, color=TRUE_BLACK)

    # Header bottom border
    pPr = hp._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "8",
        qn("w:space"): "4",
        qn("w:color"): "4A5F80",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Footer ---
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Ohanafy — The AI-Powered Operating System for Beverage")
    set_font(run, size=Pt(9), color=DARK_GREY)

    # --- Sample content so pandoc can extract styles ---
    doc.add_heading("Title", level=0)  # Uses Title style
    doc.add_heading("Heading 1", level=1)
    doc.add_heading("Heading 2", level=2)
    doc.add_heading("Heading 3", level=3)
    doc.add_paragraph("Body text paragraph.")

    # Sample table with branded styling
    table = doc.add_table(rows=2, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for cell in table.rows[0].cells:
        shading = cell._element.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "4A5F80",
        })
        cell._element.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=Pt(14), bold=True, color=TRUE_WHITE)
    # Body row
    for cell in table.rows[1].cells:
        shading = cell._element.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F4F2F0",
        })
        cell._element.get_or_add_tcPr().append(shading)

    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Cell 1"
    table.rows[1].cells[1].text = "Cell 2"

    # Re-apply header cell fonts after setting text
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=Pt(14), bold=True, color=TRUE_WHITE)

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
