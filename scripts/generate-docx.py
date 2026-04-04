#!/usr/bin/env python3
"""Generate the promotion pricing model DOCX from scratch using python-docx."""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# -- Title --
title = doc.add_heading('Ohanafy Promotion Pricing Model', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Data Model Documentation — Red Bull Tier Pricing Example').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# -- Executive Summary --
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'This document defines the data model for Ohanafy\'s promotion pricing system. '
    'The system manages on-invoice discounts and billbacks that reduce front line prices '
    'for beverage wholesalers. It is designed to be easy to manage for non-technical users, '
    'with strong analytics and reporting capabilities.'
)
doc.add_paragraph(
    'The model is demonstrated using Red Bull\'s tier pricing structure, which has 17 pricing '
    'tiers across 3 product sizes (8oz, 12oz, 16oz). Each tier defines a discount off the '
    'front line price and a billback amount the wholesaler can recoup from the supplier.'
)

# -- Glossary --
doc.add_heading('Glossary', level=1)
terms = [
    ('Front Line Price (FLP)', 'The base price before any promotions. Varies by warehouse location and retailer type. The most a retailer will ever pay. Managed by a separate system.'),
    ('On-Invoice Discount', 'A price reduction that appears directly on the retailer\'s invoice. Reduces the FLP to the "on-invoice price."'),
    ('Billback', 'An amount the wholesaler bills back to the supplier (e.g., Red Bull) to recoup promotion costs. Does NOT appear on the retailer\'s invoice.'),
    ('Promotion Program', 'A top-level container grouping related pricing tiers for a single supplier (e.g., "Red Bull 2025").'),
    ('Pricing Group', 'A named tier within a program (e.g., "Blue Diamond PCAC"). Defines discount and billback rules. Accounts and chains are assigned to pricing groups.'),
    ('Product Scope', 'Criteria defining which products a promotion covers: supplier, brand family, brand, and unit of measure (UOM).'),
    ('Net Cost', 'The true cost to the wholesaler for a promotion: Discount minus Billback.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Term'
hdr[1].text = 'Definition'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True
for term, defn in terms:
    row = table.add_row().cells
    row[0].text = term
    row[1].text = defn

doc.add_paragraph('')

# -- How Pricing Flows --
doc.add_heading('How Pricing Flows', level=1)
doc.add_paragraph('Front Line Price (from FLP system)')
doc.add_paragraph('  minus  On-Invoice Discount (from pricing group rule)')
doc.add_paragraph('  equals  On-Invoice Price (what the retailer pays)')
doc.add_paragraph('')
doc.add_paragraph('Wholesaler cost of promotion = On-Invoice Discount')
doc.add_paragraph('Wholesaler recoups from supplier = Billback')
doc.add_paragraph('Net cost to wholesaler = Discount - Billback')

# -- Data Model Overview --
doc.add_heading('Data Model Overview', level=1)
doc.add_paragraph(
    'The model has a clear hierarchy: a Promotion Program contains Product Scopes (what products) '
    'and Pricing Groups (what tiers). Each Pricing Group has Rules (discount/billback per product scope), '
    'Account/Chain Assignments (who gets this pricing), and optional Location and Retailer Group Scopes.'
)

doc.add_heading('Entity Hierarchy', level=2)
doc.add_paragraph('Promotion Program (e.g., "Red Bull 2025")')
doc.add_paragraph('    Product Scope[] (supplier + brand family + brand + UOM)', style='List Bullet')
doc.add_paragraph('    Pricing Group[] (tiers: "Blue Diamond PCAC", "Diamond PC", etc.)', style='List Bullet')
doc.add_paragraph('        Pricing Rules[] (discount + billback per product scope)', style='List Bullet 2')
doc.add_paragraph('        Account/Chain Assignments[] (with date ranges)', style='List Bullet 2')
doc.add_paragraph('        Location Scope[] (warehouse filters)', style='List Bullet 2')
doc.add_paragraph('        Retailer Group Scope[] (retailer type filters)', style='List Bullet 2')

# -- Tables --
doc.add_heading('Database Tables', level=1)

tables_info = [
    ('promotion_program', 'Top-level container. One per supplier per program period.', [
        ('id', 'UUID', 'Primary key'),
        ('name', 'VARCHAR(200)', 'Display name, e.g., "Red Bull 2025 Q3"'),
        ('supplier_id', 'UUID', 'FK to supplier'),
        ('start_date', 'DATE', 'Program start'),
        ('end_date', 'DATE', 'Program end (NULL = open-ended)'),
        ('status', 'VARCHAR(20)', 'draft, active, expired, archived'),
        ('notes', 'TEXT', 'Free-form notes'),
    ]),
    ('product_scope', 'Defines which products a program covers. Reused across all tiers.', [
        ('id', 'UUID', 'Primary key'),
        ('program_id', 'UUID', 'FK to promotion_program'),
        ('name', 'VARCHAR(100)', 'Display name, e.g., "Red Bull 8oz"'),
        ('brand_family_id', 'UUID', 'FK to brand_family (nullable)'),
        ('brand_id', 'UUID', 'FK to brand (nullable)'),
        ('product_id', 'UUID', 'FK to product (nullable)'),
        ('uom', 'VARCHAR(20)', 'e.g., "8oz", "12oz", "16oz"'),
    ]),
    ('pricing_group', 'A named tier within a program. Accounts are assigned to pricing groups.', [
        ('id', 'UUID', 'Primary key'),
        ('program_id', 'UUID', 'FK to promotion_program'),
        ('name', 'VARCHAR(100)', 'Display name, e.g., "Blue Diamond PCAC"'),
        ('tier_level', 'VARCHAR(50)', 'Normalized level: "blue_diamond", etc.'),
        ('tier_rank', 'INTEGER', 'Ordering (1 = best tier)'),
        ('account_type_modifier', 'VARCHAR(20)', '"PCAC", "PC", or NULL'),
        ('channel', 'VARCHAR(50)', '"Liquor", "Grocery", or NULL'),
        ('sort_order', 'INTEGER', 'Display ordering'),
        ('is_active', 'BOOLEAN', 'Soft disable'),
    ]),
    ('pricing_group_rule', 'Discount and billback config for a tier + product scope.', [
        ('id', 'UUID', 'Primary key'),
        ('pricing_group_id', 'UUID', 'FK to pricing_group'),
        ('product_scope_id', 'UUID', 'FK to product_scope'),
        ('discount_type', 'VARCHAR(20)', 'amount_off, percent_off, price_override'),
        ('discount_value', 'DECIMAL(10,4)', 'Discount amount/percent/price'),
        ('billback_type', 'VARCHAR(20)', 'amount, percent, override, or NULL'),
        ('billback_value', 'DECIMAL(10,4)', 'Billback value, NULL if none'),
        ('on_invoice_price', 'DECIMAL(10,4)', 'FLP - discount (stored for ops)'),
        ('front_line_price', 'DECIMAL(10,4)', 'Denormalized FLP for validation'),
    ]),
    ('pricing_group_assignment', 'Assigns accounts/chains to tiers with date ranges.', [
        ('id', 'UUID', 'Primary key'),
        ('pricing_group_id', 'UUID', 'FK to pricing_group'),
        ('account_id', 'UUID', 'FK to account (one of account/chain required)'),
        ('chain_id', 'UUID', 'FK to chain (one of account/chain required)'),
        ('effective_date', 'DATE', 'When assignment starts'),
        ('end_date', 'DATE', 'When it ends (NULL = indefinite)'),
        ('notes', 'TEXT', 'e.g., "Promoted from Gold per Q3 review"'),
    ]),
    ('pricing_group_location_scope', 'Warehouse filters. No rows = applies to all.', [
        ('id', 'UUID', 'Primary key'),
        ('pricing_group_id', 'UUID', 'FK to pricing_group'),
        ('warehouse_id', 'UUID', 'FK to warehouse'),
    ]),
    ('pricing_group_retailer_scope', 'Retailer group filters. No rows = applies to all.', [
        ('id', 'UUID', 'Primary key'),
        ('pricing_group_id', 'UUID', 'FK to pricing_group'),
        ('retailer_group_id', 'UUID', 'FK to retailer_group'),
    ]),
]

for tname, tdesc, tcols in tables_info:
    doc.add_heading(tname, level=2)
    doc.add_paragraph(tdesc)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Light Grid Accent 1'
    h = t.rows[0].cells
    h[0].text = 'Column'
    h[1].text = 'Type'
    h[2].text = 'Description'
    for cell in h:
        cell.paragraphs[0].runs[0].bold = True
    for cname, ctype, cdesc in tcols:
        row = t.add_row().cells
        row[0].text = cname
        row[1].text = ctype
        row[2].text = cdesc
    doc.add_paragraph('')

# -- Red Bull Example --
doc.add_heading('Red Bull Tier Pricing Example', level=1)
doc.add_paragraph(
    'The Red Bull 2025 promotion program has 17 pricing tiers, 3 product scopes (8oz, 12oz, 16oz), '
    'and 51 pricing rules (17 tiers x 3 UOMs). Front line prices are constant across all tiers.'
)

doc.add_heading('Front Line Prices', level=2)
flp_table = doc.add_table(rows=1, cols=2)
flp_table.style = 'Light Grid Accent 1'
flp_table.rows[0].cells[0].text = 'UOM'
flp_table.rows[0].cells[1].text = 'Front Line Price'
for cell in flp_table.rows[0].cells:
    cell.paragraphs[0].runs[0].bold = True
for uom, price in [('8oz', '$41.75'), ('12oz', '$54.45'), ('16oz', '$35.11')]:
    row = flp_table.add_row().cells
    row[0].text = uom
    row[1].text = price

doc.add_paragraph('')
doc.add_heading('Complete Tier Matrix', level=2)

tiers_data = [
    ('Blue Diamond PCAC', 1, 6.75, 6.75, 2.00, 35.00, 47.70, 33.11, 3.13, 1.25),
    ('Blue Diamond PC', 2, 6.50, 6.50, 2.00, 35.25, 47.95, 33.11, 3.00, 1.25),
    ('Blue Diamond', 3, 4.25, 4.25, 2.00, 37.50, 50.20, 33.11, 1.88, 1.25),
    ('Triple Diamond PCAC', 4, 6.50, 6.50, 2.00, 35.25, 47.95, 33.11, 2.88, 1.25),
    ('Triple Diamond PC', 5, 6.25, 6.25, 2.00, 35.50, 48.20, 33.11, 2.75, 1.25),
    ('Triple Diamond', 6, 4.00, 4.00, 2.00, 37.75, 50.45, 33.11, 1.63, 1.25),
    ('Double Diamond PC', 7, 6.00, 6.00, 2.00, 35.75, 48.45, 33.11, 2.50, 1.25),
    ('Diamond PCAC', 8, 5.75, 5.75, 2.00, 36.00, 48.70, 33.11, 2.38, 1.25),
    ('Diamond PC', 9, 5.50, 5.50, 2.00, 36.25, 48.95, 33.11, 2.25, 1.25),
    ('Diamond', 10, 3.25, 3.25, 2.00, 38.50, 51.20, 33.11, 1.13, 1.25),
    ('Platinum PC', 11, 4.25, 4.25, 2.00, 37.50, 50.20, 33.11, 1.63, 1.25),
    ('Gold PC', 12, 3.25, 3.25, 2.00, 38.50, 51.20, 33.11, 1.13, 1.25),
    ('Bronze PC', 13, 3.00, 3.00, 2.00, 38.75, 51.45, 33.11, 0, 0),
    ('LSR', 14, 3.00, 3.00, 0, 38.75, 51.45, 35.11, 0, 0),
    ('Liquor PC', 15, 4.50, 4.50, 2.00, 37.25, 49.95, 33.11, 1.75, 1.25),
    ('Grocery Diamond PC', 16, 5.50, 5.50, 2.00, 36.25, 48.95, 33.11, 2.25, 0),
    ('Grocery Plat PC', 17, 4.25, 4.25, 2.00, 37.50, 50.20, 33.11, 1.63, 0),
]

headers = ['Tier', 'Rank', 'Disc 8oz', 'Disc 12oz', 'Disc 16oz', 'Inv 8oz', 'Inv 12oz', 'Inv 16oz', 'BB 8&12oz', 'BB 16oz']
matrix = doc.add_table(rows=1, cols=len(headers))
matrix.style = 'Light Grid Accent 1'
matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers):
    matrix.rows[0].cells[i].text = h
    matrix.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    matrix.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(8)

for tier in tiers_data:
    row = matrix.add_row().cells
    row[0].text = tier[0]
    row[1].text = str(tier[1])
    for j in range(2, 10):
        val = tier[j]
        row[j].text = f'${val:.2f}' if val > 0 else '—'
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)

doc.add_paragraph('')

# -- Key Observations --
doc.add_heading('Key Observations', level=2)
observations = [
    'Discount amounts always equal FLP minus On-Invoice price (e.g., 8oz: $41.75 - $35.00 = $6.75)',
    '8oz and 12oz discounts are always identical dollar amounts',
    '16oz discount is uniformly $2.00 across most tiers (LSR has no 16oz discount)',
    'Billback for 8oz and 12oz share the same rate',
    'Some tiers have no billback (Bronze PC, LSR)',
    'Channel-specific tiers exist for Liquor and Grocery',
    'Tier names encode a level (Blue Diamond > Triple Diamond > Diamond > etc.) plus an optional account type modifier (PCAC, PC, or standard)',
]
for obs in observations:
    doc.add_paragraph(obs, style='List Bullet')

# -- Managing Tier Assignments --
doc.add_heading('Managing Tier Assignments', level=1)
doc.add_paragraph(
    'Accounts and chains are assigned to pricing groups (tiers) with date ranges. This design supports '
    'the critical workflow of moving accounts between tiers as their status changes.'
)
doc.add_heading('Moving an Account Between Tiers', level=2)
steps = [
    'Set the end_date on the current assignment to the day before the move',
    'Create a new assignment with the new tier and the effective_date of the move',
    'The old assignment is preserved in history for analytics and audit',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Precedence Rules', level=2)
doc.add_paragraph('Account-level assignment overrides chain-level assignment', style='List Bullet')
doc.add_paragraph('More specific product scope matches override broader scopes', style='List Bullet')
doc.add_paragraph('If no promotion applies, the front line price is used (no discount)', style='List Bullet')

# -- Analytics --
doc.add_heading('Analytics & Reporting', level=1)
doc.add_paragraph('The data model supports these critical business analytics:')

analytics = [
    ('Revenue Impact per Tier', 'Total discount given by each pricing group, broken down by UOM.'),
    ('Billback Recovery Rate', 'What percentage of discounts is recouped from the supplier. Critical for understanding promotion ROI.'),
    ('Net Promotion Cost', 'True cost to the wholesaler per tier after billback recovery. This is the key metric: Discount - Billback.'),
    ('Account Tier Migration History', 'Track how accounts have moved between tiers over time using the date-ranged assignments.'),
    ('Current Tier Distribution', 'How many accounts/chains are assigned to each tier right now.'),
    ('Program Comparison', 'Compare discount structures across program periods (Q1 vs Q2, 2024 vs 2025).'),
]
for name, desc in analytics:
    doc.add_heading(name, level=2)
    doc.add_paragraph(desc)

# -- FAQ --
doc.add_heading('Frequently Asked Questions', level=1)

faqs = [
    ('What happens when front line prices change?',
     'The on_invoice_price and front_line_price stored on rules may need updating. The system flags rules when FLP changes for the same product scope.'),
    ('Can an account be in two tiers at once?',
     'No. The system validates that date ranges don\'t overlap for the same account within a program. An account is always in exactly one tier (or none) at any given time.'),
    ('What if a chain is in Tier A but one of its stores needs Tier B?',
     'Create an account-level assignment for that specific store. Account-level overrides chain-level.'),
    ('How do I set up a new promotion program?',
     '1) Create the program with supplier and dates. 2) Define product scopes (which products/UOMs). 3) Create pricing groups (tiers). 4) Add rules to each tier. 5) Assign accounts/chains to tiers.'),
    ('What does "empty scope = all" mean?',
     'If a pricing group has no rows in the location scope table, it applies to all warehouses. If it has no rows in the retailer scope table, it applies to all retailer groups. You only add rows to restrict scope.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.bold = True
    doc.add_paragraph(f'A: {a}')
    doc.add_paragraph('')

# Save
out_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'docs', 'promotion-pricing-model.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
